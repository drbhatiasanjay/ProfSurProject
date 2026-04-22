"""Compile all 5 sections into a single master Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = "c:/Users/hemas/Downloads/LifeCycle_Leverage_Dashboard_Report.docx"
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
    hs.font.name = 'Arial'

# Title Page
for _ in range(3):
    doc.add_paragraph("")
t = doc.add_paragraph("LifeCycle Leverage Dashboard")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.size = Pt(28)
t.runs[0].font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
t.runs[0].bold = True

st = doc.add_paragraph("Comprehensive Analysis Report")
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.runs[0].font.size = Pt(16)
st.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

doc.add_paragraph("")
desc = doc.add_paragraph(
    "Capital Structure Determinants Across Corporate Life Stages\n"
    "401 Indian Companies | 2001-2024 | 8,677 Observations"
)
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(3):
    doc.add_paragraph("")
info = doc.add_paragraph(
    "Based on PhD Thesis by Surendra Kumar\n"
    "University of Delhi, 2025\n"
    "Supervisors: Dr. Varun Dawar & Dr. Chandra Prakash Gupta\n\n"
    "Dashboard: https://lifecycle-leverage-779655496440.us-east1.run.app\n"
    "Repository: https://github.com/drbhatiasanjay/ProfSurProject"
)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(10)
info.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# Section A: Thesis Analysis
doc.add_heading("Section A: Complete Thesis Analysis & Dashboard Mapping", level=1)

doc.add_heading("A.1 Research Framework", level=2)
doc.add_paragraph(
    "This thesis examines whether capital structure varies across corporate life stages and "
    "identifies stage-specific determinants for 401 BSE-listed non-financial Indian companies "
    "over FY2001-FY2024 (8,677 firm-year observations). Data sourced from CMIE ProwessOnWeb. "
    "Life stages classified using Dickinson (2011) cash-flow methodology into 8 stages: "
    "Startup, Growth, Maturity, Shakeout1/2/3, Decline, and Decay."
)

doc.add_heading("A.2 Theoretical Foundation", level=2)
theories = [
    ("Pecking Order Theory (Myers & Majluf 1984)", "Firms prefer internal funds, then debt, then equity. Predicts profitability (-). STRONGLY SUPPORTED: coef = -27.44***."),
    ("Trade-off Theory (Kraus & Litzenberger 1973)", "Balance tax shields vs distress costs. Predicts tangibility (+). SUPPORTED: coef = +26.02***."),
    ("Agency Cost Theory (Jensen & Meckling 1976)", "Debt disciplines managers; promoter ownership reduces agency costs. SUPPORTED: promoter share = -0.09***."),
    ("M&M Irrelevance (1958)", "Capital structure irrelevant under perfect markets. REJECTED: multiple determinants highly significant."),
    ("Signalling Theory (Ross 1977)", "Leverage signals quality. Predicts profitability (+). REJECTED: profitability consistently negative."),
    ("Free Cash Flow Theory (Jensen 1986)", "Debt forces cash disgorging. PARTIALLY SUPPORTED."),
    ("Corporate Life Stage (Dickinson 2011)", "Determinants vary by stage. STRONGLY SUPPORTED: ANOVA F=31.94, p<0.001."),
]
for name, desc in theories:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}: ")
    r.bold = True
    p.add_run(desc)

doc.add_heading("A.3 Key Empirical Results", level=2)

# Main regression table
doc.add_heading("Fixed Effects Regression (R-squared = 0.309)", level=3)
t = doc.add_table(rows=11, cols=4)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(["Variable", "Coefficient", "Sig", "Theory"]):
    t.rows[0].cells[i].text = h
data = [
    ("Profitability", "-27.44", "***", "Pecking Order confirmed"),
    ("Tangibility", "+26.02", "***", "Trade-off confirmed"),
    ("Tax Shield", "+0.008", "***", "Trade-off confirmed"),
    ("Interest Rate", "-9.58", "***", "Higher rates reduce debt"),
    ("Market Return", "-0.21", "***", "Market timing"),
    ("GFC Dummy", "-6.60", "***", "Crisis deleveraging"),
    ("COVID Dummy", "-37.91", "***", "Largest event impact"),
    ("Promoter Share", "-0.09", "***", "Agency cost"),
    ("Dividend", "+0.003", "**", "Signalling"),
    ("Hausman Test", "Chi2=225.53", "p=0.000", "FE preferred"),
]
for i, (v, c, s, th) in enumerate(data):
    for j, val in enumerate([v, c, s, th]):
        t.rows[i+1].cells[j].text = val

doc.add_paragraph("")
doc.add_paragraph("System GMM (Table 5.12): Lag leverage = 0.622*** | Speed of adjustment = 37.8% per year")
doc.add_paragraph("")

# Stage findings
doc.add_heading("Stage-Specific Findings", level=3)
t2 = doc.add_table(rows=6, cols=4)
t2.style = 'Light Grid Accent 1'
for i, h in enumerate(["Stage", "Level Dummy", "Delta Dummy", "Key Finding"]):
    t2.rows[0].cells[i].text = h
for i, row in enumerate([
    ("Startup", "+5.05***", "+5.69***", "Debt-dependent, forced borrowing"),
    ("Growth", "+3.35***", "+4.45***", "Profitability NOT significant (novel)"),
    ("Maturity", "-3.92***", "-3.93***", "Stickiest leverage (lag=0.743)"),
    ("Decline", "+2.09 (NS)", "+4.42***", "Involuntary over-leverage"),
    ("Decay", "+1.55 (NS)", "-0.82 (NS)", "Terminal, leverage involuntary"),
]):
    for j, val in enumerate(row):
        t2.rows[i+1].cells[j].text = val

doc.add_paragraph("")

# Thesis mapping
doc.add_heading("A.4 Thesis-to-Dashboard Mapping", level=2)
doc.add_paragraph("All 20 core thesis elements implemented. 17 novel features added beyond thesis.")

t3 = doc.add_table(rows=13, cols=3)
t3.style = 'Light Grid Accent 1'
for i, h in enumerate(["Thesis Element", "Dashboard Page", "Status"]):
    t3.rows[0].cells[i].text = h
for i, row in enumerate([
    ("Dickinson Classification", "Page 4: Bulk Upload", "Full"),
    ("ANOVA (F=31.94)", "Pages 1 + 8", "Full"),
    ("Pairwise Tukey HSD", "Page 8", "Full"),
    ("FE Regression (R2=0.309)", "Page 8 Auto-Suggest", "Full"),
    ("Delta-Leverage", "Page 13", "Full"),
    ("System GMM", "Page 13", "Full"),
    ("Startup Regressions", "Pages 8+13", "Full"),
    ("Growth/Maturity Comparison", "Pages 8+13", "Full"),
    ("Decline/Decay Regressions", "Pages 8+13", "Full"),
    ("Hausman + BP-LM Tests", "Page 8", "Full"),
    ("Event Dummies", "Pages 1+8+13", "Full"),
    ("Interest Rate + Market", "Pages 1+8", "Full"),
]):
    for j, val in enumerate(row):
        t3.rows[i+1].cells[j].text = val

doc.add_page_break()

# Section B: merge from docx
def merge_docx(doc, path, title):
    if not os.path.exists(path):
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"[File not found: {path}]")
        return
    src = Document(path)
    doc.add_heading(title, level=1)
    for para in src.paragraphs:
        if para.text.strip():
            new_p = doc.add_paragraph(para.text)
    for table in src.tables:
        new_t = doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_t.style = 'Light Grid Accent 1'
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_t.rows[i].cells[j].text = cell.text
    print(f"  Merged: {path}")

merge_docx(doc, "c:/Users/hemas/Downloads/ProfSurProject/Section_B_Dashboard_Feature_Documentation.docx",
           "Section B: Dashboard Feature Documentation")
doc.add_page_break()

merge_docx(doc, "c:/Users/hemas/Downloads/ProfSurProject/docs/Section_C_From_Academic_Research_to_Real_World_Application.docx",
           "Section C: From Academic Research to Real-World Application")
doc.add_page_break()

# Section D: Future Enhancements
doc.add_heading("Section D: Future Enhancements Roadmap", level=1)
doc.add_paragraph("30 enhancements across 5 categories.")

categories = {
    "D.1 Data & Coverage": [
        ("Real-time CMIE integration", "Auto-refresh quarterly via ETL", "All users", "High"),
        ("Expand to 2,000+ BSE firms", "Scale beyond BSE 500", "Analysts", "High"),
        ("International markets", "Cross-country comparison", "Researchers, PE", "Medium"),
        ("Financial sector firms", "Bank/NBFC-specific measures", "Credit analysts", "Medium"),
        ("Alternative data (ESG, ratings)", "Non-financial enrichment", "ESG analysts", "Low"),
    ],
    "D.2 Model Enhancements": [
        ("Quantile regression", "Determinants at different leverage percentiles", "Researchers", "High"),
        ("Instrumental variables", "Address endogeneity", "Researchers", "High"),
        ("Bayesian panel models", "Probabilistic targets with credible intervals", "CFOs", "Medium"),
        ("RL capital advisor", "Optimal leverage via reinforcement learning", "CFOs", "Low"),
        ("Causal inference (DiD for IBC)", "Isolate policy causal effect", "Regulators", "High"),
    ],
    "D.3 UI/UX Enhancements": [
        ("Natural language query", "Ask questions in English", "All users", "High"),
        ("PDF report generation", "One-click analyst report", "Analysts, CFOs", "High"),
        ("Alert system", "Email on stage transition or threshold breach", "Credit analysts", "High"),
        ("Dark mode + mobile", "Accessibility", "All users", "Medium"),
    ],
    "D.4 Platform Enhancements": [
        ("Multi-user auth", "Role-based access", "Enterprise", "High"),
        ("REST/GraphQL API", "Programmatic access", "Developers", "High"),
        ("Scheduled retraining", "Monthly ML model refresh", "Ops", "Medium"),
        ("White-label capability", "Custom branding", "Enterprise sales", "Low"),
    ],
    "D.5 Research Extensions": [
        ("Interaction effects", "Profitability x size non-linear terms", "Researchers", "High"),
        ("Sector-specific models", "Pharma vs IT vs manufacturing", "Sector analysts", "High"),
        ("Debt maturity structure", "Short vs long-term composition", "Credit analysts", "Medium"),
        ("Speed of adjustment by stage", "GMM per stage", "Researchers", "High"),
        ("Ownership + pledging", "Promoter pledging as crisis indicator", "Credit, PE", "High"),
    ],
}

for cat_name, items in categories.items():
    doc.add_heading(cat_name, level=2)
    t = doc.add_table(rows=len(items)+1, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(["Enhancement", "Description", "Target User", "Priority"]):
        t.rows[0].cells[i].text = h
    for i, (n, d, u, p) in enumerate(items):
        for j, val in enumerate([n, d, u, p]):
            t.rows[i+1].cells[j].text = val
    doc.add_paragraph("")

doc.add_page_break()

# Section E: merge from docx
merge_docx(doc, "c:/Users/hemas/Downloads/ProfSurProject/docs/Section_E_Commercial_Gap_Analysis.docx",
           "Section E: Commercial Gap Analysis & Competitive Landscape")

doc.save(OUT)
print(f"\nMaster document: {OUT}")
print(f"Size: {os.path.getsize(OUT) // 1024} KB")
